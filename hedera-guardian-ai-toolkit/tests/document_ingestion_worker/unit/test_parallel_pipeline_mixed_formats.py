"""
Unit tests for mixed format support in parallel_pipeline.py.

Test organization:
- TestMixedFormatDiscovery: Document discovery for PDF and DOCX files
- TestFormatBasedRouting: Parser routing based on file extension
- TestUnsupportedFormatHandling: Behavior with unsupported file types
- TestSourceFormatMetadata: source_format field in pipeline state
"""

from pathlib import Path

import pytest
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_test_pdf(output_path: Path, content: str, title: str = "Test Document"):
    """
    Create a simple PDF file for testing.

    Args:
        output_path: Path where PDF should be saved
        content: Text content for the PDF
        title: Title of the document
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter

    # Add title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, title)

    # Add content
    c.setFont("Helvetica", 12)
    y_position = height - 100
    lines = content.split("\n")

    for line in lines:
        if y_position < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y_position = height - 50

        c.drawString(50, y_position, line)
        y_position -= 20

    c.save()


def create_test_docx(output_path: Path, content: str, title: str = "Test Document"):
    """
    Create a simple DOCX file for testing.

    Args:
        output_path: Path where DOCX should be saved
        content: Text content for the document
        title: Title of the document
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in content.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(output_path)


@pytest.fixture
def sample_docx_file(sample_pdf_dir):
    """Create a sample DOCX file for testing."""
    docx_path = sample_pdf_dir / "test_document.docx"
    create_test_docx(docx_path, "Sample DOCX content for testing", "Test Document")
    return docx_path


@pytest.fixture
def mixed_documents_dir(tmp_path):
    """Create a directory with mixed PDF and DOCX files."""
    doc_dir = tmp_path / "mixed_docs"
    doc_dir.mkdir()

    # Create PDF files
    create_test_pdf(doc_dir / "document1.pdf", "PDF content 1", "PDF 1")
    create_test_pdf(doc_dir / "document2.pdf", "PDF content 2", "PDF 2")

    # Create DOCX files
    create_test_docx(doc_dir / "document3.docx", "DOCX content 3", "DOCX 3")
    create_test_docx(doc_dir / "document4.docx", "DOCX content 4", "DOCX 4")

    # Create unsupported files
    (doc_dir / "document5.doc").write_text("Legacy Word content")
    (doc_dir / "document6.rtf").write_text("RTF content")
    (doc_dir / "document7.txt").write_text("Text content")

    return doc_dir


class TestMixedFormatDiscovery:
    """Test document discovery for PDF and DOCX files."""

    def test_discover_pdf_files_only(self, tmp_path):
        """Test that PDF files are discovered."""
        doc_dir = tmp_path / "docs"
        doc_dir.mkdir()
        create_test_pdf(doc_dir / "test1.pdf", "Content 1")
        create_test_pdf(doc_dir / "test2.pdf", "Content 2")

        # Discover files
        pdf_files = list(doc_dir.glob("*.pdf"))
        assert len(pdf_files) == 2

    def test_discover_docx_files_only(self, tmp_path):
        """Test that DOCX files are discovered."""
        doc_dir = tmp_path / "docs"
        doc_dir.mkdir()
        create_test_docx(doc_dir / "test1.docx", "Content 1")
        create_test_docx(doc_dir / "test2.docx", "Content 2")

        # Discover files
        docx_files = list(doc_dir.glob("*.docx"))
        assert len(docx_files) == 2

    def test_discover_mixed_formats(self, mixed_documents_dir):
        """Test that both PDF and DOCX files are discovered together."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
        )

        # Discover all supported files
        all_files = []
        for pattern in DocumentParserFactory.get_supported_glob_patterns():
            all_files.extend(mixed_documents_dir.glob(pattern))

        # Should find PDFs and DOCXs only
        assert len(all_files) == 4
        extensions = {f.suffix.lower() for f in all_files}
        assert extensions == {".pdf", ".docx"}

    def test_discover_excludes_unsupported_formats(self, mixed_documents_dir):
        """Test that unsupported formats are excluded from discovery."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
        )

        # Discover supported files
        supported_files = []
        for pattern in DocumentParserFactory.get_supported_glob_patterns():
            supported_files.extend(mixed_documents_dir.glob(pattern))

        # Should not include .doc, .rtf, .txt
        extensions = {f.suffix.lower() for f in supported_files}
        assert ".doc" not in extensions
        assert ".rtf" not in extensions
        assert ".txt" not in extensions


class TestFormatBasedRouting:
    """Test parser routing based on file extension."""

    def test_get_format_for_pdf(self):
        """Test format detection for PDF files."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
            SupportedFormat,
        )

        result = DocumentParserFactory.get_format(Path("document.pdf"))
        assert result == SupportedFormat.PDF

    def test_get_format_for_docx(self):
        """Test format detection for DOCX files."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
            SupportedFormat,
        )

        result = DocumentParserFactory.get_format(Path("document.docx"))
        assert result == SupportedFormat.DOCX

    def test_routing_selects_correct_parser(self):
        """Test that format routing would select correct parser type."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
            SupportedFormat,
        )

        # Simulate routing logic
        files = [
            Path("doc1.pdf"),
            Path("doc2.docx"),
            Path("doc3.PDF"),
            Path("doc4.DOCX"),
        ]

        routing = {}
        for f in files:
            fmt = DocumentParserFactory.get_format(f)
            if fmt:
                routing[f.name] = fmt

        assert routing["doc1.pdf"] == SupportedFormat.PDF
        assert routing["doc2.docx"] == SupportedFormat.DOCX
        assert routing["doc3.PDF"] == SupportedFormat.PDF
        assert routing["doc4.DOCX"] == SupportedFormat.DOCX


class TestUnsupportedFormatHandling:
    """Test behavior with unsupported file types."""

    def test_unsupported_format_returns_none(self):
        """Test that unsupported formats return None from factory."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
        )

        unsupported = [
            Path("document.doc"),  # Legacy Word
            Path("document.rtf"),
            Path("document.txt"),
            Path("document.odt"),
            Path("document.xlsx"),
        ]

        for f in unsupported:
            assert DocumentParserFactory.get_format(f) is None
            assert DocumentParserFactory.is_supported(f) is False

    def test_filtering_unsupported_files(self, mixed_documents_dir):
        """Test filtering a directory to exclude unsupported files."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
        )

        all_files = list(mixed_documents_dir.iterdir())
        supported_files = [f for f in all_files if DocumentParserFactory.is_supported(f)]
        unsupported_files = [f for f in all_files if not DocumentParserFactory.is_supported(f)]

        # 4 supported (2 PDF + 2 DOCX)
        assert len(supported_files) == 4

        # 3 unsupported (.doc, .rtf, .txt)
        assert len(unsupported_files) == 3


class TestSourceFormatMetadata:
    """Test source_format metadata in pipeline state."""

    def test_source_format_derived_from_pdf_extension(self):
        """Test that source_format is derived correctly for PDF."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
            SupportedFormat,
        )

        file_path = Path("/documents/methodology.pdf")
        fmt = DocumentParserFactory.get_format(file_path)

        assert fmt == SupportedFormat.PDF
        assert fmt.value == "pdf"

    def test_source_format_derived_from_docx_extension(self):
        """Test that source_format is derived correctly for DOCX."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            DocumentParserFactory,
            SupportedFormat,
        )

        file_path = Path("/documents/methodology.docx")
        fmt = DocumentParserFactory.get_format(file_path)

        assert fmt == SupportedFormat.DOCX
        assert fmt.value == "docx"

    def test_source_format_value_is_string(self):
        """Test that source_format values are suitable for JSON/Qdrant storage."""
        from document_ingestion_worker.document_parsing.document_parser_factory import (
            SupportedFormat,
        )

        # SupportedFormat is a str enum, values should be strings
        assert isinstance(SupportedFormat.PDF.value, str)
        assert isinstance(SupportedFormat.DOCX.value, str)

        # Direct string comparison should work
        assert SupportedFormat.PDF == "pdf"
        assert SupportedFormat.DOCX == "docx"


class TestSingleDocumentStateSourceFormat:
    """Test source_format field in SingleDocumentState."""

    def test_single_document_state_accepts_source_format(self):
        """Test that SingleDocumentState can include source_format field."""
        from document_ingestion_worker.models import SingleDocumentState

        # Create state with source_format
        state: SingleDocumentState = {
            "document_id": "test_doc",
            "pdf_path": Path("test.docx"),
            "staged_path": Path("staged/test_doc"),
            "source_format": "docx",
            "start_from": "beginning",
        }

        assert state["source_format"] == "docx"
        assert state["pdf_path"].suffix == ".docx"

    def test_create_single_document_state_without_source_format(self):
        """Test that existing create_single_document_state works (no source_format yet)."""
        from document_ingestion_worker.models import create_single_document_state

        # Current function doesn't include source_format - this tests backward compatibility
        state = create_single_document_state(
            pdf_path=Path("test.pdf"),
            staged_path=Path("staged/test"),
        )

        assert state["document_id"] == "test"
        assert state["pdf_path"] == Path("test.pdf")
        # source_format may not be present in current implementation
        # This test ensures we don't break existing code
